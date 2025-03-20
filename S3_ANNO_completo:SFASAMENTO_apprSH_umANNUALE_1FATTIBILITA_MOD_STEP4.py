import pandas as pd
from openpyxl import Workbook
import numpy as np
import time
from docx import Document

# Inizializzo il tempo di inizio per misurare il tempo computazionale della simulazione
start_time = time.time()

# Defininsco la soglia unica per lo step 3,4
single_thr = 0.50

## PERCORSI DEI FILE DI INPUT E OUTPUT ##

# Input
input_path = '/Users/giulianosmarrazzo/Desktop/TESI_COCA_COLA/FILE PER PYTHON/SCENARI/3/INPUT.xlsx'
statistiche_path = '/Users/giulianosmarrazzo/Desktop/TESI_COCA_COLA/FILE PER PYTHON/SCENARI/3/LOGICA_STEP3.xlsx'

# Output
output_shipm_in_out_prod_path = f'/Users/giulianosmarrazzo/Desktop/TESI_COCA_COLA/OUTPUT PYTHON/Simulazioni/SCENARI/SCENARIO 3/SIMULAZIONI/ANNO/{int(single_thr*100)}/sped_ingressi_uscite_prod.xlsx'
output_stock_path = f'/Users/giulianosmarrazzo/Desktop/TESI_COCA_COLA/OUTPUT PYTHON/Simulazioni/SCENARI/SCENARIO 3/SIMULAZIONI/ANNO/{int(single_thr*100)}/scorte.xlsx'
output_shuttling_and_cap_path = f'/Users/giulianosmarrazzo/Desktop/TESI_COCA_COLA/OUTPUT PYTHON/Simulazioni/SCENARI/SCENARIO 3/SIMULAZIONI/ANNO/{int(single_thr*100)}/shuttling_e_cap.xlsx'
output_shuttling_summary_path = f'/Users/giulianosmarrazzo/Desktop/TESI_COCA_COLA/OUTPUT PYTHON/Simulazioni/SCENARI/SCENARIO 3/SIMULAZIONI/ANNO/{int(single_thr*100)}/info_shuttling.xlsx'
output_step4_case2_path = f'/Users/giulianosmarrazzo/Desktop/TESI_COCA_COLA/OUTPUT PYTHON/Simulazioni/SCENARI/SCENARIO 3/SIMULAZIONI/ANNO/{int(single_thr*100)}/risultati_step4_caso2.xlsx'
output_word_path = f'/Users/giulianosmarrazzo/Desktop/TESI_COCA_COLA/OUTPUT PYTHON/Simulazioni/SCENARI/SCENARIO 3/SIMULAZIONI/ANNO/{int(single_thr*100)}/report.docx'

# INIZIALIZZAZIONE FILE EXCEL PER output_step4_case2_path

wb = Workbook()
wb.save(output_step4_case2_path)

## CARICAMENTO DEL FILE DI INPUT E LETTURA DEI FOGLI ##

excel_data = pd.ExcelFile(input_path)
in_shipm_dataset = excel_data.parse('SPED_INGRESSO')
out_shipm_dataset = excel_data.parse('SPED_USCITA')
in_dataset = excel_data.parse('IN_ASIS')
out_dataset = excel_data.parse('OUT_ASIS')
prod_dataset = excel_data.parse('PROD6372_ARR_ECC')
initial_stock_dataset = excel_data.parse('GiacenzeIniziali_neg_sfas_ECC')
item_r_s_dataset = excel_data.parse('R-S')
item_fam_dataset = excel_data.parse('ITEM-FAMIGLIA')
stat_dataset72 = pd.read_excel(statistiche_path, sheet_name='STATISTICHE_72')
stat_dataset71 = pd.read_excel(statistiche_path, sheet_name='STATISTICHE_71')
stat_dataset18 = pd.read_excel(statistiche_path, sheet_name='STATISTICHE_18')

## VARIABILI DI SUPPORTO ALLA SIMULAZIONE ##

# Lista dei materiali e delle date
item_list = initial_stock_dataset['ITEM'].tolist()
date_list = [col.strftime('%d/%m/%Y') for col in pd.to_datetime(prod_dataset.columns[1:], dayfirst=True)]

# Dizionario che associa a ciascun item la propria uscita media #####
mean_out_dict72 = dict(zip(stat_dataset72['ITEM'], stat_dataset72['USCITA MEDIA']))
mean_out_dict71 = dict(zip(stat_dataset71['ITEM'], stat_dataset71['USCITA MEDIA']))
mean_out_dict18 = dict(zip(stat_dataset18['ITEM'], stat_dataset18['USCITA MEDIA']))

# Elenco completo dei magazzini e di quelli di interesse
all_warehouses = [6372, 6385, 6318, 6371]
warehouses = [6372, 6371, 6318]

# Caratterizzazione capacità originali e aggiuntive dei magazzini e dell' hub
original_cap_72 = 15000
original_cap_71 = 24000
original_cap_18 = 1500
# In S4 avevamo original_cap_72 = 15000, original_cap_71 = 23000 e additional_cap_71 = 7645, quindi hub_cap = 45645.
# Ora abbiamo original_cap_72 = 15000, original_cap_71 = 24000 e original_cap_18 = 1500, quindi
# additional_cap_71 = 45645 - 15000 - 24000 - 1500 = 5145
additional_cap_71 = 5145
final_cap_71 = original_cap_71 + additional_cap_71
warehouses_cap_dict = {'6372': original_cap_72, '6371': final_cap_71, '6318': original_cap_18}
hub_cap = sum(warehouses_cap_dict[str(warehouse)] for warehouse in warehouses)

# Ricreo il DF della prod_plt per convertire le date nel formato corretto e inserire tutti gli item
# Converto le colonne di prod_dataset in formato stringa 'gg/mm/aaaa'
prod_dataset.columns = ['ITEM'] + date_list
complete_prod_df = pd.DataFrame(0.0, index=item_list, columns=date_list)
for _, row in prod_dataset.iterrows():
    item = row['ITEM']
    if item in item_list:
        complete_prod_df.loc[item, row.index.intersection(date_list)] = row[row.index.intersection(date_list)].fillna(0)
complete_prod_df.insert(0, 'ITEM', complete_prod_df.index)

# Dizionario che associa l' item alla famiglia
item_fam_dict = dict(zip(item_fam_dataset['Item'], item_fam_dataset['Famiglia']))

# Filtra i runner e gli stranger
runners = [item for item in item_list if item in item_r_s_dataset[item_r_s_dataset['R-S'] == 'R']['ITEM'].values]
strangers = [item for item in item_list if item in item_r_s_dataset[item_r_s_dataset['R-S'] == 'S']['ITEM'].values]

# Filtra gli item per famiglie (runner AA e A e stranger B, C e CC)
items_r_A = [item for item, fam in item_fam_dict.items() if fam.endswith('A') and not fam.endswith('AA')]
items_r_AA = [item for item, fam in item_fam_dict.items() if fam.endswith('AA')]
items_s_B = [item for item, fam in item_fam_dict.items() if fam.endswith('B')]
items_s_C = [item for item, fam in item_fam_dict.items() if fam.endswith('C') and not fam.endswith('CC')]
items_s_CC = [item for item, fam in item_fam_dict.items() if fam.endswith('CC')]

# # Inizializzazione delle soglia minime e massime di bilanciamento (capire se e come utilizzarle) #####
min_balancing_thr_A = 0.20
min_balancing_thr_AA = 0.20
min_balancing_thr_B = 0
min_balancing_thr_C = 0
min_balancing_thr_CC = 0
max_balancing_thr_A = single_thr
max_balancing_thr_AA = single_thr
max_balancing_thr_B = 0
max_balancing_thr_C = 0
max_balancing_thr_CC = 0

# Creo dei contatori per monitorare le occorrenze nello step 2
step2_case1_count = 0
step2_case2p1_count = 0
step2_case2p2_count = 0
step2_case2p3_count = 0
step2_case3p1_count = 0
step2_case3p2_count = 0
step2_case3p3_count = 0
step2_case4p1_count = 0
step2_case4p2_count = 0
step2_case4p3_count = 0
step2_case5_count = 0
step2_case6_count = 0
step2_case7_count = 0
step2_case8_count = 0
difference_double_cap_shutt_count = 0

# Creo dei contatori per monitorare la corretteza dello step 3,4
no_overcap_count = 0
overcap_72_count = 0
overcap_71_count = 0
overcap_18_count = 0
overcap_72_71_count = 0
overcap_72_18_count = 0
overcap_71_18_count = 0
overcap_72_71_18_count = 0

# Contatori per caso 2 step 3,4
case2_only_cap_shutt_71_count = 0
case2_only_cap_shutt_18_count = 0
case2_double_cap_shutt_count = 0

# Definisco lo step_size per la determinazione di cov_day_star nello step 4 (capire se e come utilizzarle) #####
step_size = 0.1

# Creo un contatore per verificare che il fabbisogno sia correttamente soddsfattO in ogni giorno
needs_check_passed = 0
needs_check_failed = 0
cap_check_passed = 0
cap_check_failed = 0

## CREAZIONE DEI DF DI OUTPUT ##

# Creo dei DF per tener traccia del numero di spedizioni inziali, riallocate e finali
in_shipm_dict_df = {warehouse: pd.DataFrame(0.0, index=['INIZIALI', 'RIALLOCATE+', 'RIALLOCATE-', 'FINALI'], columns=date_list) for warehouse in warehouses}
out_shipm_dict_df = {warehouse: pd.DataFrame(0.0, index=['INIZIALI', 'RIALLOCATE+', 'RIALLOCATE-', 'FINALI'], columns=date_list) for warehouse in warehouses}
all_shipm_dict_df = {warehouse: pd.DataFrame(0.0, index=['INIZIALI', 'RIALLOCATE+', 'RIALLOCATE-', 'FINALI'], columns=date_list) for warehouse in warehouses}

# Creo ora due dizionari (uno per le spedizioni in ingresso e uno per quelle in uscita) in cui la chiave principale è il
# magazzino, quella secondaria è la data di date_list ed il valore è una lista contenente i codici delle spedizioni in quella
# data e per quel magazzino
in_shipm_code_dict = {warehouse: {date: [] for date in date_list} for warehouse in warehouses}
out_shipm_code_dict = {warehouse: {date: [] for date in date_list} for warehouse in warehouses}

# Creo i DF per gli ingressi e le uscite relative ai magazzini
in_item_dict_df = {warehouse: pd.DataFrame(0.0, index=item_list, columns=date_list) for warehouse in warehouses}
out_item_dict_df = {warehouse: pd.DataFrame(0.0, index=item_list, columns=date_list) for warehouse in warehouses}

# Creo i DF per gli shuttling da fabbisogno di materiale (needs) e quelli per il rispetto della capacità (cap)
# Struttura nidificata: shutt_type[partenza][destinazione]
# Creiamo un elenco con le combinazioni dei magazzini
needs_shutt_dict_df = {
    src: {  # Primo livello: chiave è il magazzino di partenza (src)
        dest: pd.DataFrame(0.0, index=item_list, columns=date_list)
        # Secondo livello: chiave è il magazzino di destinazione (dest)
        for dest in warehouses if dest != src  # Evita di includere il magazzino stesso come destinazione
    } for src in warehouses  # Ciclo su tutti i magazzini per creare il livello esterno
}
cap_shutt_dict_df = {
    src: {
        dest: pd.DataFrame(0.0, index=item_list, columns=date_list)
        for dest in warehouses if dest != src
    } for src in warehouses
}
total_shutt_dict_df = {
    src: {
        dest: pd.DataFrame(0.0, index=item_list, columns=date_list)
        for dest in warehouses if dest != src
    } for src in warehouses
}
warehouse_combinations = [f"{src}-{dest}" for src in warehouses for dest in warehouses if src != dest]
total_cap_shutt_df = pd.DataFrame(0.0, index=warehouse_combinations + ["Totali"], columns=date_list)
total_shutt_hub_df = pd.DataFrame(0.0, index=item_list, columns=date_list)

#  Creo i DF per le scorte provvisorie (temporary), quelli in seguito agli shuttling da fabbisogno di materiale (needs)
#  e infine quelli aggiornati in seguito agli shuttling per il rispetto della capacità (final)
temporary_stock_dict_df = {warehouse: pd.DataFrame(0.0, index=item_list, columns=date_list) for warehouse in warehouses}
needs_shutt_stock_dict_df = {warehouse: pd.DataFrame(0.0, index=item_list, columns=date_list) for warehouse in warehouses}
final_stock_dict_df = {warehouse: pd.DataFrame(0.0, index=item_list, columns=date_list) for warehouse in warehouses}
final_stock_hub_df = pd.DataFrame(0.0, index=item_list, columns=date_list)
# Creo un DF per tenere traccia della scorta aggiuntiva necessaria dovuta ad errori di approssimazioni
additional_stock_df = pd.DataFrame(0.0, index=item_list, columns=["SCORTA AGGIUNTIVA"])

# Creo i DF per tenere traccia delle eccedenze per ciascun magazzino
overcap_dict_df = {warehouse: pd.DataFrame(0.0, index=['OVERCAP'], columns=date_list) for warehouse in warehouses}

# Creo il df per tenere traccia delle soglie di bilanciamento finali
balancing_thr_df = pd.DataFrame(0.0, index=['SOGLIA FINALE A', 'SOGLIA FINALE AA', 'SOGLIA FINALE B', 'SOGLIA FINALE C', 'SOGLIA FINALE CC'], columns=date_list)

# Creo il DF per memorizzare gli shuttling totali effettivi per il caso 2 dello step 4
total_eff_shutt_case2 = pd.DataFrame(0.0, index=['SHUTT EFFETTIVI', '6372-6371', '6372-6318'], columns=date_list)

# Intestazioni colonne dei DF
for warehouse in warehouses:
    in_shipm_dict_df[warehouse].index.name = 'SPEDIZIONI'
    out_shipm_dict_df[warehouse].index.name = 'SPEDIZIONI'
    all_shipm_dict_df[warehouse].index.name = 'SPEDIZIONI'
    in_item_dict_df[warehouse].index.name = 'ITEM'
    out_item_dict_df[warehouse].index.name = 'ITEM'
    temporary_stock_dict_df[warehouse].index.name = 'ITEM'
    needs_shutt_stock_dict_df[warehouse].index.name = 'ITEM'
    final_stock_dict_df[warehouse].index.name = 'ITEM'
total_cap_shutt_df.index.name = 'SH_CAP_TOTALI'
total_shutt_hub_df.index.name = 'ITEM'
final_stock_hub_df.index.name = 'ITEM'
additional_stock_df.index.name = 'ITEM'
for src in warehouses:
    for dest in warehouses:
        if src != dest:
            # Shuttling da fabbisogno
            needs_shutt_dict_df[src][dest].index.name = 'ITEM'
            # Shuttling da capacità
            cap_shutt_dict_df[src][dest].index.name = 'ITEM'
            # Shuttling totali
            total_shutt_dict_df[src][dest].index.name = 'ITEM'

## SIMULAZIONE ##

for current_date_idx, current_date in enumerate(date_list):

    # Memorizzo il giorno successivo (per assicurare lo sfasamento)
    next_date = date_list[current_date_idx + 1] if current_date_idx + 1 < len(date_list) else None

    # STEP 1: ALLOCAZIONE SPEDIZIONI AI MAGAZZINI
    # In questo step si effettua una riallocazione dinamica delle spedizioni in ingresso ed in uscita secondo le logiche
    # previste dallo scenario che consente di tenere traccia della variaizione nel numero di spedizioni gestite da ciascun magazzino

    # Gestione spedizioni in ingresso:
    daily_in_shipm = in_shipm_dataset[in_shipm_dataset['Posting Date'] == pd.to_datetime(current_date, dayfirst=True)]

    # Conteggio delle spedizioni iniziali con Plant == 6372
    in_shipm_dict_df[6372].at['INIZIALI', current_date] = daily_in_shipm[daily_in_shipm['Plant'] == 6372].shape[0]

    # Conteggio delle spedizioni iniziali con Plant == 6371
    in_shipm_dict_df[6371].at['INIZIALI', current_date] = daily_in_shipm[daily_in_shipm['Plant'] == 6371].shape[0]

    # Conteggio delle spedizioni iniziali con Plant == 6318
    in_shipm_dict_df[6318].at['INIZIALI', current_date] = daily_in_shipm[daily_in_shipm['Plant'] == 6318].shape[0]

    # Riallocazione e conteggio delle spedizioni per popolare di DF in_shipm_code_dict e in_shipm_dict_df
    for index, row in daily_in_shipm.iterrows():
        if row['FR'] != 0 and row['FS'] == 0 and row['MR'] == 0 and row['MS'] == 0:
            in_shipm_code_dict[6372][current_date].append(row['Reference'])
            if row['Plant'] == 6371:
                in_shipm_dict_df[6372].at['RIALLOCATE+', current_date] += 1
                in_shipm_dict_df[6371].at['RIALLOCATE-', current_date] += 1
            elif row['Plant'] == 6318:
                in_shipm_dict_df[6372].at['RIALLOCATE+', current_date] += 1
                in_shipm_dict_df[6318].at['RIALLOCATE-', current_date] += 1
            elif row['Plant'] in [6385]:
                in_shipm_dict_df[6372].at['RIALLOCATE+', current_date] += 1
        elif row['Plant'] == 6318:
            in_shipm_code_dict[6318][current_date].append(row['Reference'])
        else:
            in_shipm_code_dict[6371][current_date].append(row['Reference'])
            if row['Plant'] == 6372:
                in_shipm_dict_df[6371].at['RIALLOCATE+', current_date] += 1
                in_shipm_dict_df[6372].at['RIALLOCATE-', current_date] += 1
            elif row['Plant'] in [6385]:
                in_shipm_dict_df[6371].at['RIALLOCATE+', current_date] += 1

    # Gestione spedizioni in uscita:
    daily_out_shipm = out_shipm_dataset[out_shipm_dataset['CurrLoadSt'] == pd.to_datetime(current_date, dayfirst=True)]

    # Conteggio delle spedizioni iniziali con ShPt == 6372
    out_shipm_dict_df[6372].at['INIZIALI', current_date] = daily_out_shipm[daily_out_shipm['ShPt'] == 6372].shape[0]

    # Conteggio delle spedizioni iniziali con ShPt == 6371
    out_shipm_dict_df[6371].at['INIZIALI', current_date] = daily_out_shipm[daily_out_shipm['ShPt'] == 6371].shape[0]

    # Conteggio delle spedizioni iniziali con ShPt == 6318
    out_shipm_dict_df[6318].at['INIZIALI', current_date] = daily_out_shipm[daily_out_shipm['ShPt'] == 6318].shape[0]

    # Riallocazione e conteggio delle spedizioni per popolare i DF ref_sped_uscita_dict e out_shipm_dict_df
    for index, row in daily_out_shipm.iterrows():
        if row['Tipologia Carico'] == 'FTL' and row['FR'] != 0 and row['FS'] == 0 and row['MR'] == 0 and row['MS'] == 0:
            out_shipm_code_dict[6372][current_date].append(row['Shipment'])
            if row['ShPt'] == 6371:
                out_shipm_dict_df[6372].at['RIALLOCATE+', current_date] += 1
                out_shipm_dict_df[6371].at['RIALLOCATE-', current_date] += 1
            elif row['ShPt'] == 6318:
                out_shipm_dict_df[6372].at['RIALLOCATE+', current_date] += 1
                out_shipm_dict_df[6318].at['RIALLOCATE-', current_date] += 1
            elif row['ShPt'] in [6385]:
                out_shipm_dict_df[6372].at['RIALLOCATE+', current_date] += 1
        elif row['Tipologia Carico'] == 'LTL':
            out_shipm_code_dict[6318][current_date].append(row['Shipment'])
            if row['ShPt'] == 6372:
                out_shipm_dict_df[6318].at['RIALLOCATE+', current_date] += 1
                out_shipm_dict_df[6372].at['RIALLOCATE-', current_date] += 1
            elif row['ShPt'] == 6371:
                out_shipm_dict_df[6318].at['RIALLOCATE+', current_date] += 1
                out_shipm_dict_df[6371].at['RIALLOCATE-', current_date] += 1
            elif row['ShPt'] in [6385]:
                out_shipm_dict_df[6318].at['RIALLOCATE+', current_date] += 1
        else:
            out_shipm_code_dict[6371][current_date].append(row['Shipment'])
            if row['ShPt'] == 6372:
                out_shipm_dict_df[6371].at['RIALLOCATE+', current_date] += 1
                out_shipm_dict_df[6372].at['RIALLOCATE-', current_date] += 1
            elif row['ShPt'] == 6318:
                out_shipm_dict_df[6371].at['RIALLOCATE+', current_date] += 1
                out_shipm_dict_df[6318].at['RIALLOCATE-', current_date] += 1
            elif row['ShPt'] in [6385]:
                out_shipm_dict_df[6371].at['RIALLOCATE+', current_date] += 1

    if next_date:
        # Gestione spedizioni in uscita il giorno successivo (utile per applicare lo sfasamento):
        next_date_out_shipm = out_shipm_dataset[
            out_shipm_dataset['CurrLoadSt'] == pd.to_datetime(next_date, dayfirst=True)]
        for index, row in next_date_out_shipm.iterrows():
            if row['Tipologia Carico'] == 'FTL' and row['FR'] != 0 and row['FS'] == 0 and row['MR'] == 0 and row[
                'MS'] == 0:
                out_shipm_code_dict[6372][next_date].append(row['Shipment'])
            elif row['Tipologia Carico'] == 'LTL':
                out_shipm_code_dict[6318][next_date].append(row['Shipment'])
            else:
                out_shipm_code_dict[6371][next_date].append(row['Shipment'])

    for warehouse in warehouses:
        # Calcolo spedizioni finali in ingresso e in uscita
        in_shipm_dict_df[warehouse].at['FINALI', current_date] = (
                in_shipm_dict_df[warehouse].at['INIZIALI', current_date] +
                in_shipm_dict_df[warehouse].at['RIALLOCATE+', current_date] -
                in_shipm_dict_df[warehouse].at['RIALLOCATE-', current_date]
        )

        out_shipm_dict_df[warehouse].at['FINALI', current_date] = (
                out_shipm_dict_df[warehouse].at['INIZIALI', current_date] +
                out_shipm_dict_df[warehouse].at['RIALLOCATE+', current_date] -
                out_shipm_dict_df[warehouse].at['RIALLOCATE-', current_date]
        )

        # Calcolo spedizioni totali
        all_shipm_dict_df[warehouse].at['INIZIALI', current_date] = (
                in_shipm_dict_df[warehouse].at['INIZIALI', current_date] +
                out_shipm_dict_df[warehouse].at['INIZIALI', current_date]
        )

        all_shipm_dict_df[warehouse].at['RIALLOCATE+', current_date] = (
                in_shipm_dict_df[warehouse].at['RIALLOCATE+', current_date] +
                out_shipm_dict_df[warehouse].at['RIALLOCATE+', current_date]
        )

        all_shipm_dict_df[warehouse].at['RIALLOCATE-', current_date] = (
                in_shipm_dict_df[warehouse].at['RIALLOCATE-', current_date] +
                out_shipm_dict_df[warehouse].at['RIALLOCATE-', current_date]
        )

        all_shipm_dict_df[warehouse].at['FINALI', current_date] = (
                in_shipm_dict_df[warehouse].at['FINALI', current_date] +
                out_shipm_dict_df[warehouse].at['FINALI', current_date]
        )

    ## STEP 2: CALCOLO DEGLI SHUTTLING NECESSARI (DA FABBISOGNO) ##

    # In questo step si procede a determinare il numero di plt per ogni item in ingresso/uscita da a/da ciascun magazzino e,
    # nota anche la produzione, al calcolo delle scorte provvisorie (al netto degli shuttling).
    # Quindi si passa alla determinazione degli shuttling da fabbisogno (al fine di assicurare che in ciascun giorno,
    # si disponga per ciascun item, dello stock necessario a soddisfare le uscite del giorno successivo).
    # Infine si procede all' aggiornamento degli shuttling da fabbisogno e alla determinazione delle scorte in seguito
    # a tali shuttling.

    # Calcolo ingressi e uscite per item e per magazzino alla data corrente
    for warehouse in warehouses:

        # Calcolo plt in ingresso di tutti gli item alla data corrente
        if in_shipm_code_dict[warehouse][current_date]:  # Se in_shipm_code non è vuota
            filtered_in = in_dataset[in_dataset['Reference'].isin(in_shipm_code_dict[warehouse][current_date])]
            # Raggruppa il DataFrame filtrato per 'Material' e calcola la somma dei valori nella colonna 'plt' per ciascun materiale
            in_group = filtered_in.groupby('Material')['plt'].sum()
            for item, total_plt in in_group.items():
                if item in in_item_dict_df[warehouse].index:
                    in_item_dict_df[warehouse].at[item, current_date] = total_plt

        # Calcolo plt in uscita di tutti gli item alla data corrente
        if out_shipm_code_dict[warehouse][current_date]:
            filtered_out = out_dataset[out_dataset['Shipment'].isin(out_shipm_code_dict[warehouse][current_date])]
            out_group = filtered_out.groupby('Material')['plt'].sum()
            for item, total_plt in out_group.items():
                if item in out_item_dict_df[warehouse].index:
                    out_item_dict_df[warehouse].at[item, current_date] = total_plt

        # Calcolo plt in uscita di tutti gli item per il giorno successivo
        if next_date:
            if out_shipm_code_dict[warehouse][next_date]:
                next_date_filtered_out = out_dataset[
                    out_dataset['Shipment'].isin(out_shipm_code_dict[warehouse][next_date])]
                next_date_out_group = next_date_filtered_out.groupby('Material')['plt'].sum()
                for item, total_plt in next_date_out_group.items():
                    if item in out_item_dict_df[warehouse].index:
                        out_item_dict_df[warehouse].at[item, next_date] = total_plt

    # Calcolo delle scorte provvisorie dei diversi item alla data corrente
    for item in item_list:
        for warehouse in warehouses:
            # Calcola la giacenza iniziale di un materiale specifico in un magazzino per il primo giorno (current_date_idx == 0).
            # Cerca il valore corrispondente nel DataFrame 'initial_stock_dataset' nella colonna del magazzino ('Giacenza_{warehouse}').
            # Se il materiale non è presente nel DataFrame, assegna la giacenza iniziale a 0.
            if current_date_idx == 0:  # Primo Giorno
                initial_stock_plt = initial_stock_dataset.loc[initial_stock_dataset['ITEM'] == item, f'Giacenza_{warehouse}'].values[0] if item in initial_stock_dataset['ITEM'].values else 0.0
                in_plt = in_item_dict_df[warehouse].at[item, current_date]
                out_plt = out_item_dict_df[warehouse].at[item, current_date]

                if warehouse == 6372:
                    prod_plt = complete_prod_df.at[item, current_date] if item in complete_prod_df.index else 0.0
                    temporary_stock_dict_df[warehouse].at[item, current_date] = (
                            initial_stock_plt + in_plt + prod_plt - out_plt
                    )
                else:  # Magazzino 6371 e 6318 (no prod_plt)
                    temporary_stock_dict_df[warehouse].at[item, current_date] = (
                            initial_stock_plt + in_plt - out_plt
                    )

            else:  # Giorni successivi
                prev_day_stock = final_stock_dict_df[warehouse].at[item, date_list[current_date_idx - 1]]
                in_plt = in_item_dict_df[warehouse].at[item, current_date]
                out_plt = out_item_dict_df[warehouse].at[item, current_date]

                if warehouse == 6372:
                    prod_plt = complete_prod_df.at[item, current_date] if item in complete_prod_df.index else 0.0
                    temporary_stock_dict_df[warehouse].at[item, current_date] = (
                            prev_day_stock + in_plt + prod_plt - out_plt
                    )
                else:  # Magazzino 6371 e 6318 (no prod_plt)
                    temporary_stock_dict_df[warehouse].at[item, current_date] = (
                            prev_day_stock + in_plt - out_plt
                    )

            # Inizializzo lo stock dopo gli shuttling da fabbisogno di ciascun item e per ciascun magazzino alla rispettiva scorta provvisoria
            # per facilitare il successivo aggiornamento
            needs_shutt_stock_dict_df[warehouse].at[item, current_date] = temporary_stock_dict_df[warehouse].at[item, current_date]

        # CRITERIO SHUTTLING DA FABBISOGNO
        # Privilegiamo sempre lo shuttling da un unico magazzino ed in ogni caso facciamo in modo  da ridurre
        # al minimo gli shutling da fabbisogno da 72 verso gli altri magazzini al fine di prevenire backshuttling.

        # Logica degli shuttling da fabbisogno finc al 30/12/3023 (per assicurare lo sfasamento)

        if next_date:

            # CASO 1: Scorta provvisoria di ciascun magazzino >= delle rispettiva uscita del giorno successivo
            if (temporary_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date] >= 0 and
                temporary_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date] >= 0 and
                temporary_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date] >= 0):
                needs_shutt_stock_dict_df[6372].at[item, current_date] += 0
                needs_shutt_stock_dict_df[6371].at[item, current_date] += 0
                needs_shutt_stock_dict_df[6318].at[item, current_date] += 0
                step2_case1_count += 1

            # CASI NON UNIVOCAMENTE DETERMINATI

            # CASO 2: Scorta provvisoria di 72 e 71 >= delle riuspettive uscite del giorno successivo e
            # Scorta provvisoria di 18 < della rispettiv uscita del giorno successivo
            elif (temporary_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date] >= 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date] >= 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date] < 0):
                needs_shutt_amount = out_item_dict_df[6318].at[item, next_date] - temporary_stock_dict_df[6318].at[item, current_date]
                surplus72 = temporary_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date]
                surplus71 = temporary_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date]
                # 2.1: il surplus di 71 consente di soddisfare interamente il fabbisogno di 18, allora in tal caso si effettua
                # solo lo shuttling da 71 a 18, così da evitare di spostare item da 72 e di effettuare shuttling in ingresso
                # da entrambi i magazzini
                if surplus71 >= needs_shutt_amount:
                    step2_case2p1_count += 1
                    needs_shutt_dict_df[6371][6318].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6371].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6318].at[item, current_date] += needs_shutt_amount
                # 2.2: il surplus di 72 consente di soddisfare interamente il fabbisogno di 18, allora in tal caso si effettua
                # solo lo shuttling da 72 a 18, così da evitare di effettuare shuttling in ingresso da entrambi i magazzini
                elif surplus72 >= np.ceil(needs_shutt_amount):
                    step2_case2p2_count += 1
                    needs_shutt_amount = np.ceil(needs_shutt_amount)
                    needs_shutt_dict_df[6372][6318].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6372].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6318].at[item, current_date] += needs_shutt_amount
                # 2.3: sfruttiamo prima il surplus di 71 e la restante parte viene compensata da 72
                else:
                    step2_case2p3_count += 1
                    shutt_71_18 = surplus71
                    needs_shutt_dict_df[6371][6318].at[item, current_date] = shutt_71_18
                    needs_shutt_stock_dict_df[6371].at[item, current_date] -= shutt_71_18
                    shutt_72_18 = np.ceil(needs_shutt_amount - shutt_71_18)
                    needs_shutt_dict_df[6372][6318].at[item, current_date] = shutt_72_18
                    needs_shutt_stock_dict_df[6372].at[item, current_date] -= shutt_72_18
                    needs_shutt_stock_dict_df[6318].at[item, current_date] += shutt_71_18 + shutt_72_18

            # CASO 3: Scorta provvisoria di 72 e 18 >= delle riuspettive uscite del giorno successivo e
            # Scorta provvisoria di 71 < della rispettive uscita del giorno successivo
            elif (temporary_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date] >= 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date] < 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date] >= 0):
                needs_shutt_amount = out_item_dict_df[6371].at[item, next_date] - temporary_stock_dict_df[6371].at[item, current_date]
                surplus72 = temporary_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date]
                surplus18 = temporary_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date]
                # 3.1: il surplus di 72 consente di soddisfare interamente il fabbisogno di 71, allora in tal caso si effettua
                # solo lo shuttling da 72 a 71, così da evitare di effettuare shuttling in ingresso da entrambi i magazzini
                if surplus72 >= np.ceil(needs_shutt_amount):
                    step2_case3p2_count += 1
                    needs_shutt_amount = np.ceil(needs_shutt_amount)
                    needs_shutt_dict_df[6372][6371].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6372].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6371].at[item, current_date] += needs_shutt_amount
                # 3.2: il surplus di 18 consente di soddisfare interamente il fabbisogno di 71, allora in tal caso si effettua
                # solo lo shuttling da 18 a 71, così da evitare effettuare shuttling in ingresso da entrambi i magazzini
                elif surplus18 >= needs_shutt_amount:
                    step2_case3p1_count += 1
                    needs_shutt_dict_df[6318][6371].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6318].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6371].at[item, current_date] += needs_shutt_amount
                # 3.3: sfruttiamo prima il surplus di 18 e la restante parte viene compensata da 72
                else:
                    step2_case3p3_count += 1
                    shutt_72_71 = surplus72
                    needs_shutt_dict_df[6372][6371].at[item, current_date] = shutt_72_71
                    needs_shutt_stock_dict_df[6372].at[item, current_date] -= shutt_72_71
                    shutt_18_71 = needs_shutt_amount - shutt_72_71
                    needs_shutt_dict_df[6318][6371].at[item, current_date] = shutt_18_71
                    needs_shutt_stock_dict_df[6318].at[item, current_date] -= shutt_18_71
                    needs_shutt_stock_dict_df[6371].at[item, current_date] += shutt_72_71 + shutt_18_71

            # CASO 4: Scorta provvisoria di 71 e 18 >= delle riuspettive uscite del giorno successivo e
            # Scorta provvisoria di 72 < della rispettiv uscita del giorno successivo
            elif (temporary_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date] < 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date] >= 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date] >= 0):
                needs_shutt_amount = out_item_dict_df[6372].at[item, next_date] - temporary_stock_dict_df[6372].at[item, current_date]
                surplus71 = temporary_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date]
                surplus18 = temporary_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date]
                # 4.1: il surplus di 71 consente di soddisfare interamente il fabbisogno di 72, allora in tal caso si effettua
                # solo lo shuttling da 71 a 72, così da evitare di effettuare shuttling in ingresso da entrambi i magazzini
                # Perchè si preferisce 71 a 18?
                # Il magazzino 18 ha una capacità molto limitata e, se anche fosse possibilie, non varrebbe la pena utilizarlo per soddisfare il
                # fabbisogno di 72 (si rischierebbe solo di dover effettuare nuovi shuttling da fabbisogno verso 18).
                if surplus71 >= needs_shutt_amount:
                    step2_case4p1_count += 1
                    needs_shutt_dict_df[6371][6372].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6371].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6372].at[item, current_date] += needs_shutt_amount
                # 4.2: il surplus di 18 consente di soddisfare interamente il fabbisogno di 72, allora in tal caso si effettua
                # solo lo shuttling da 18 a 72, così da evitare di effettuare shuttling in ingresso da entrambi i magazzini
                elif surplus18 >= needs_shutt_amount:
                    step2_case4p2_count += 1
                    needs_shutt_dict_df[6318][6372].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6318].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6372].at[item, current_date] += needs_shutt_amount
                # 4.3: sfruttiamo prima il surplus di 71 e la restante parte viene compensata da 18
                else:
                    step2_case4p3_count += 1
                    shutt_71_72 = np.floor(surplus71)
                    needs_shutt_dict_df[6371][6372].at[item, current_date] = shutt_71_72
                    needs_shutt_stock_dict_df[6371].at[item, current_date] -= shutt_71_72
                    shutt_18_72 = needs_shutt_amount - shutt_71_72
                    if shutt_71_72 + shutt_18_72 < needs_shutt_amount:
                        print(f" Scorte insufficienti in 6318 a coprire le uscite del giorno successivo di 6372 "
                              f"per l' ITEM {item}, alla data {current_date}:")
                    needs_shutt_dict_df[6318][6372].at[item, current_date] = shutt_18_72
                    needs_shutt_stock_dict_df[6318].at[item, current_date] -= shutt_18_72
                    needs_shutt_stock_dict_df[6372].at[item, current_date] += shutt_71_72 + shutt_18_72

            # CASI UNIVOCAMENTE DETERMINATI

            # CASO 5: Scorta provvisoria di 71 e 18 < delle rispettive uscite del giorno successivo
            elif (temporary_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date] >= 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date] < 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date] < 0):
                step2_case5_count += 1
                shutt_72_71 = np.ceil(out_item_dict_df[6371].at[item, next_date] - temporary_stock_dict_df[6371].at[item, current_date])
                shutt_72_18 = np.ceil(out_item_dict_df[6318].at[item, next_date] - temporary_stock_dict_df[6318].at[item, current_date])
                needs_shutt_dict_df[6372][6371].at[item, current_date] = shutt_72_71
                needs_shutt_dict_df[6372][6318].at[item, current_date] = shutt_72_18
                needs_shutt_stock_dict_df[6372].at[item, current_date] -= shutt_72_71 + shutt_72_18
                needs_shutt_stock_dict_df[6371].at[item, current_date] += shutt_72_71
                needs_shutt_stock_dict_df[6318].at[item, current_date] += shutt_72_18

            # CASO 6: Scorta provvisoria di 72 e 18 < delle rispettive uscite del giorno successivo
            elif (temporary_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date] < 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date] >= 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date] < 0):
                step2_case6_count += 1
                shutt_71_72 = out_item_dict_df[6372].at[item, next_date] - temporary_stock_dict_df[6372].at[item, current_date]
                shutt_71_18 = out_item_dict_df[6318].at[item, next_date] - temporary_stock_dict_df[6318].at[item, current_date]
                needs_shutt_dict_df[6371][6372].at[item, current_date] = shutt_71_72
                needs_shutt_dict_df[6371][6318].at[item, current_date] = shutt_71_18
                needs_shutt_stock_dict_df[6371].at[item, current_date] -= shutt_71_72 + shutt_71_18
                needs_shutt_stock_dict_df[6372].at[item, current_date] += shutt_71_72
                needs_shutt_stock_dict_df[6318].at[item, current_date] += shutt_71_18

            # CASO 7: Scorta provvisoria di 72 e 71 < delle rispettive uscite del giorno successivo
            elif (temporary_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date] < 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date] < 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date] >= 0):
                step2_case7_count += 1
                shutt_18_72 = out_item_dict_df[6372].at[item, next_date] - temporary_stock_dict_df[6372].at[item, current_date]
                shutt_18_71 = out_item_dict_df[6371].at[item, next_date] - temporary_stock_dict_df[6371].at[item, current_date]
                needs_shutt_dict_df[6318][6372].at[item, current_date] = shutt_18_72
                needs_shutt_dict_df[6318][6371].at[item, current_date] = shutt_18_71
                needs_shutt_stock_dict_df[6318].at[item, current_date] -= shutt_18_72 + shutt_18_71
                needs_shutt_stock_dict_df[6372].at[item, current_date] += shutt_18_72
                needs_shutt_stock_dict_df[6371].at[item, current_date] += shutt_18_71

            # CASO 8: Scorta provvisoria di 72, 71 e 18 < delle rispettive uscite del giorno successivo
            elif (temporary_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date] < 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date] < 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date] < 0):
                step2_case8_count += 1
                # CASO 4: Scorte provvisorie per entrambi i magazzini al giorno corrente < uscite da entrambi i magazzini giorno successivo
                print(f" Scorte insufficienti a coprire le uscite del giorno successivo per ITEM {item} in tutti i magazzini alla data {current_date}:")
                missing72 = out_item_dict_df[6372].at[item, next_date] - temporary_stock_dict_df[6372].at[item, current_date]
                missing71 = out_item_dict_df[6371].at[item, next_date] - temporary_stock_dict_df[6371].at[item, current_date]
                missing18 = out_item_dict_df[6318].at[item, next_date] - temporary_stock_dict_df[6318].at[item, current_date]
                print(f"Scorte mancanti per ITEM {item} alla data {current_date}:")
                print(f"   - Magazzino 6372: {missing72:.2e}")
                print(f"   - Magazzino 6371: {missing71:.2e}")
                print(f"   - Magazzino 6318: {missing18:.2e}")
                print("Correggo le scorte provvisorie e di conseguenza quelle dopo gli shuttling da fabbisogno")
                needs_shutt_stock_dict_df[6372].at[item, current_date] += missing72
                needs_shutt_stock_dict_df[6371].at[item, current_date] += missing71
                needs_shutt_stock_dict_df[6318].at[item, current_date] += missing18

        else:

            # Logica degli shuttling da fabbisogno per il 31/12/3023

            # CASO 1: Scorta provvisoria di ciascun magazzino >= delle rispettiva uscita del giorno successivo
            if (temporary_stock_dict_df[6372].at[item, current_date] >= 0 and
                temporary_stock_dict_df[6371].at[item, current_date] >= 0 and
                temporary_stock_dict_df[6318].at[item, current_date] >= 0):
                needs_shutt_stock_dict_df[6372].at[item, current_date] += 0
                needs_shutt_stock_dict_df[6371].at[item, current_date] += 0
                needs_shutt_stock_dict_df[6318].at[item, current_date] += 0
                step2_case1_count += 1

            # CASI NON UNIVOCAMENTE DETERMINATI

            # CASO 2: Scorta provvisoria di 72 e 71 >= 0 e Scorta provvisoria di 18 < 0
            elif (temporary_stock_dict_df[6372].at[item, current_date] >= 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] >= 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] < 0):
                needs_shutt_amount = -temporary_stock_dict_df[6318].at[item, current_date]
                surplus72 = temporary_stock_dict_df[6372].at[item, current_date]
                surplus71 = temporary_stock_dict_df[6371].at[item, current_date]
                # 2.1: il surplus di 71 consente di soddisfare interamente il fabbisogno di 18, allora in tal caso si effettua
                # solo lo shuttling da 71 a 18, così da evitare di spostare item da 72 e di effettuare shuttling in ingresso
                # da entrambi i magazzini
                if surplus71 >= needs_shutt_amount:
                    step2_case2p1_count += 1
                    needs_shutt_dict_df[6371][6318].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6371].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6318].at[item, current_date] += needs_shutt_amount
                # 2.2: il surplus di 72 consente di soddisfare interamente il fabbisogno di 18, allora in tal caso si effettua
                # solo lo shuttling da 72 a 18, così da evitare di effettuare shuttling in ingresso da entrambi i magazzini
                elif surplus72 >= np.ceil(needs_shutt_amount):
                    step2_case2p2_count += 1
                    needs_shutt_amount = np.ceil(needs_shutt_amount)
                    needs_shutt_dict_df[6372][6318].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6372].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6318].at[item, current_date] += needs_shutt_amount
                # 2.3: sfruttiamo prima il surplus di 71 e la restante parte viene compensata da 72
                else:
                    step2_case2p3_count += 1
                    shutt_71_18 = surplus71
                    needs_shutt_dict_df[6371][6318].at[item, current_date] = shutt_71_18
                    needs_shutt_stock_dict_df[6371].at[item, current_date] -= shutt_71_18
                    shutt_72_18 = np.ceil(needs_shutt_amount - shutt_71_18)
                    needs_shutt_dict_df[6372][6318].at[item, current_date] = shutt_72_18
                    needs_shutt_stock_dict_df[6372].at[item, current_date] -= shutt_72_18
                    needs_shutt_stock_dict_df[6318].at[item, current_date] += shutt_71_18 + shutt_72_18

            # CASO 3: Scorta provvisoria di 72 e 18 >= 0 e Scorta provvisoria di 71 < 0
            elif (temporary_stock_dict_df[6372].at[item, current_date] >= 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] < 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] >= 0):
                needs_shutt_amount = -temporary_stock_dict_df[6371].at[item, current_date]
                surplus72 = temporary_stock_dict_df[6372].at[item, current_date]
                surplus18 = temporary_stock_dict_df[6318].at[item, current_date]
                # 3.1: il surplus di 72 consente di soddisfare interamente il fabbisogno di 71, allora in tal caso si effettua
                # solo lo shuttling da 72 a 71, così da evitare di effettuare shuttling in ingresso da entrambi i magazzini
                if surplus72 >= np.ceil(needs_shutt_amount):
                    step2_case3p2_count += 1
                    needs_shutt_amount = np.ceil(needs_shutt_amount)
                    needs_shutt_dict_df[6372][6371].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6372].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6371].at[item, current_date] += needs_shutt_amount
                # 3.2: il surplus di 18 consente di soddisfare interamente il fabbisogno di 71, allora in tal caso si effettua
                # solo lo shuttling da 18 a 71, così da evitare di effettuare shuttling in ingresso da entrambi i magazzini
                elif surplus18 >= needs_shutt_amount:
                    step2_case3p1_count += 1
                    needs_shutt_dict_df[6318][6371].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6318].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6371].at[item, current_date] += needs_shutt_amount
                # 3.3: sfruttiamo prima il surplus di 18 e la restante parte viene compensata da 72
                else:
                    step2_case3p3_count += 1
                    shutt_72_71 = surplus72
                    needs_shutt_dict_df[6372][6371].at[item, current_date] = shutt_72_71
                    needs_shutt_stock_dict_df[6372].at[item, current_date] -= shutt_72_71
                    shutt_18_71 = needs_shutt_amount - shutt_72_71
                    needs_shutt_dict_df[6318][6371].at[item, current_date] = shutt_18_71
                    needs_shutt_stock_dict_df[6318].at[item, current_date] -= shutt_18_71
                    needs_shutt_stock_dict_df[6371].at[item, current_date] += shutt_72_71 + shutt_18_71

            # CASO 4: Scorta provvisoria di 71 e 18 >= 0 e Scorta provvisoria di 72 < 0
            elif (temporary_stock_dict_df[6372].at[item, current_date] < 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] >= 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] >= 0):
                needs_shutt_amount = -temporary_stock_dict_df[6372].at[item, current_date]
                surplus71 = temporary_stock_dict_df[6371].at[item, current_date]
                surplus18 = temporary_stock_dict_df[6318].at[item, current_date]
                # 4.1: il surplus di 71 consente di soddisfare interamente il fabbisogno di 72, allora in tal caso si effettua
                # solo lo shuttling da 71 a 72, così da evitare di effettuare shuttling in ingresso da entrambi i magazzini
                if surplus71 >= needs_shutt_amount:
                    step2_case4p1_count += 1
                    needs_shutt_dict_df[6371][6372].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6371].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6372].at[item, current_date] += needs_shutt_amount
                # 4.2: il surplus di 18 consente di soddisfare interamente il fabbisogno di 72, allora in tal caso si effettua
                # solo lo shuttling da 18 a 72, così da evitare di effettuare shuttling in ingresso da entrambi i magazzini
                elif surplus18 >= needs_shutt_amount:
                    step2_case4p2_count += 1
                    needs_shutt_dict_df[6318][6372].at[item, current_date] = needs_shutt_amount
                    needs_shutt_stock_dict_df[6318].at[item, current_date] -= needs_shutt_amount
                    needs_shutt_stock_dict_df[6372].at[item, current_date] += needs_shutt_amount
                # 4.3: sfruttiamo prima il surplus di 71 e la restante parte viene compensata da 18
                else:
                    step2_case4p3_count += 1
                    shutt_71_72 = np.floor(surplus71)
                    needs_shutt_dict_df[6371][6372].at[item, current_date] = shutt_71_72
                    needs_shutt_stock_dict_df[6371].at[item, current_date] -= shutt_71_72
                    shutt_18_72 = needs_shutt_amount - shutt_71_72
                    if shutt_71_72 + shutt_18_72 < needs_shutt_amount:
                        print(f" Scorte insufficienti in 6318 a coprire il fabbisogno di 6372 "
                              f"per l' ITEM {item}, alla data {current_date}:")
                    needs_shutt_dict_df[6318][6372].at[item, current_date] = shutt_18_72
                    needs_shutt_stock_dict_df[6318].at[item, current_date] -= shutt_18_72
                    needs_shutt_stock_dict_df[6372].at[item, current_date] += shutt_71_72 + shutt_18_72

            # CASI UNIVOCAMENTE DETERMINATI

            # CASO 5: Scorta provvisoria di 71 e 18 < 0
            elif (temporary_stock_dict_df[6372].at[item, current_date] >= 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] < 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] < 0):
                step2_case5_count += 1
                shutt_72_71 = np.ceil(-temporary_stock_dict_df[6371].at[item, current_date])
                shutt_72_18 = np.ceil(-temporary_stock_dict_df[6318].at[item, current_date])
                needs_shutt_dict_df[6372][6371].at[item, current_date] = shutt_72_71
                needs_shutt_dict_df[6372][6318].at[item, current_date] = shutt_72_18
                needs_shutt_stock_dict_df[6372].at[item, current_date] -= shutt_72_71 + shutt_72_18
                needs_shutt_stock_dict_df[6371].at[item, current_date] += shutt_72_71
                needs_shutt_stock_dict_df[6318].at[item, current_date] += shutt_72_18

            # CASO 6: Scorta provvisoria di 72 e 18 < 0
            elif (temporary_stock_dict_df[6372].at[item, current_date] < 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] >= 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] < 0):
                step2_case6_count += 1
                shutt_71_72 = -temporary_stock_dict_df[6372].at[item, current_date]
                shutt_71_18 = -temporary_stock_dict_df[6318].at[item, current_date]
                needs_shutt_dict_df[6371][6372].at[item, current_date] = shutt_71_72
                needs_shutt_dict_df[6371][6318].at[item, current_date] = shutt_71_18
                needs_shutt_stock_dict_df[6371].at[item, current_date] -= shutt_71_72 + shutt_71_18
                needs_shutt_stock_dict_df[6372].at[item, current_date] += shutt_71_72
                needs_shutt_stock_dict_df[6318].at[item, current_date] += shutt_71_18

            # CASO 7: Scorta provvisoria di 72 e 71 < 0
            elif (temporary_stock_dict_df[6372].at[item, current_date] < 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] < 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] >= 0):
                step2_case7_count += 1
                shutt_18_72 = -temporary_stock_dict_df[6372].at[item, current_date]
                shutt_18_71 = -temporary_stock_dict_df[6371].at[item, current_date]
                needs_shutt_dict_df[6318][6372].at[item, current_date] = shutt_18_72
                needs_shutt_dict_df[6318][6371].at[item, current_date] = shutt_18_71
                needs_shutt_stock_dict_df[6318].at[item, current_date] -= shutt_18_72 + shutt_18_71
                needs_shutt_stock_dict_df[6372].at[item, current_date] += shutt_18_72
                needs_shutt_stock_dict_df[6371].at[item, current_date] += shutt_18_71

            # CASO 8: Scorta provvisoria di 72, 71 e 18 < 0
            elif (temporary_stock_dict_df[6372].at[item, current_date] < 0 and
                  temporary_stock_dict_df[6371].at[item, current_date] < 0 and
                  temporary_stock_dict_df[6318].at[item, current_date] < 0):
                step2_case8_count += 1
                # CASO 4: Scorte provvisorie per entrambi i magazzini al giorno corrente < uscite da entrambi i magazzini giorno successivo
                print(f" Scorte negative per ITEM {item} in tutti i magazzini alla data {current_date}")
                missing72 = -temporary_stock_dict_df[6372].at[item, current_date]
                missing71 = -temporary_stock_dict_df[6371].at[item, current_date]
                missing18 = -temporary_stock_dict_df[6318].at[item, current_date]
                print(f"Scorte mancanti per ITEM {item} alla data {current_date}:")
                print(f"   - Magazzino 6372: {missing72:.2e}")
                print(f"   - Magazzino 6371: {missing71:.2e}")
                print(f"   - Magazzino 6318: {missing18:.2e}")
                print("Correggo le scorte provvisorie e di conseguenza quelle dopo gli shuttling da fabbisogno")
                needs_shutt_stock_dict_df[6372].at[item, current_date] += missing72
                needs_shutt_stock_dict_df[6371].at[item, current_date] += missing71
                needs_shutt_stock_dict_df[6318].at[item, current_date] += missing18

        # Verifico che alla data corrente tutti e 3 i magazzini abbiano una scorta in seguito agli shuttling da fabbisofno
        # iun grado di soddisfare il rispettivo fabbisogno.
        # Alla fine del codice posso usare lo stesso blocco per effettuare la stessa verififca ma in relazione alle scorte finali e in seguito
        # agli shuttling da capacità
        for warehouse in warehouses:
            if next_date:
                if needs_shutt_stock_dict_df[warehouse].at[item, current_date] - out_item_dict_df[warehouse].at[
                    item, next_date] >= 0:
                    needs_check_passed += 1
                else:
                    needs_check_failed += 1
                    additional_stock = out_item_dict_df[warehouse].at[item, next_date] - needs_shutt_stock_dict_df[warehouse].at[item, current_date]
                    print(f" Fabbisogno non correttamente soddisfatto alla data {current_date} per l' item {item}"
                          f" nel magazzino {warehouse} per una quantità pari a {additional_stock}, quindi correggo la scorta"
                          f" e aggiorno il dataframe della scorta aggiuntiva.")
                    needs_shutt_stock_dict_df[warehouse].at[item, current_date] += additional_stock
                    additional_stock_df.at[item, "SCORTA AGGIUNTIVA"] += additional_stock
            else:
                if needs_shutt_stock_dict_df[warehouse].at[item, current_date] >= 0:
                    needs_check_passed += 1
                else:
                    additional_stock = - needs_shutt_stock_dict_df[warehouse].at[item, current_date]
                    print(f" Fabbisogno non correttamente soddisfatto alla data {current_date} per l' item {item}"
                          f" nel magazzino {warehouse} per una quantità pari a {additional_stock}, quindi correggo la scorta"
                          f" e aggiorno il dataframe della scorta aggiuntiva.")
                    needs_shutt_stock_dict_df[warehouse].at[item, current_date] += additional_stock
                    additional_stock_df.at[item, "SCORTA AGGIUNTIVA"] += additional_stock

    ## STEP 3,4 - CALCOLO OVERCAPACITY e SHUTTLING PER RISPETTO CAPACITA' ##

    # In questo step si procede a calcolare l' overcapacity dei due magazzini, al fine di determinare l' ammontare degli shttling
    # da capacità da un magazzino all' altro.
    # In particolare per come sono state settate le capacità dei magazzini, gli unici casi possibili saranno 1 e 2.
    # Nel secondo caso si distingueranno le due fasi che caratterizzeranno i parametri alpha e cov_day_tar al fine di determinare
    # quali item spostare ed in che quantità sposarli.
    # Infine si procede all' aggiornamento dei relativi shuttling da capacità e totali e delle scorte finali.

    # Determinazione iniziale delle overcap dei singoli magazzini
    overcap_dict_df[6372].at['OVERCAP', current_date] = max(0, needs_shutt_stock_dict_df[6372][current_date].sum() - warehouses_cap_dict['6372'])
    overcap_dict_df[6371].at['OVERCAP', current_date] = max(0, needs_shutt_stock_dict_df[6371][current_date].sum() - warehouses_cap_dict['6371'])
    overcap_dict_df[6318].at['OVERCAP', current_date] = max(0, needs_shutt_stock_dict_df[6318][current_date].sum() - warehouses_cap_dict['6318'])

    # Aggiorno di default gli shuttling totali (per ora quelli di capacità sono tutti nulli, dunque gli shuttling totali saranno
    # proprio pari a quelli di fabbisogno) e a questi si sovrascriverà il risultato definitivo in seguito al calcolo dello shuttling
    # da capacità se necessario; analogamente per le scorte finali.
    for item in item_list:
        for src in warehouses:
            for dest in warehouses:
                if src != dest:
                    total_shutt_dict_df[src][dest].at[item, current_date] = needs_shutt_dict_df[src][dest].at[item, current_date]
                    total_shutt_hub_df.at[item, current_date] += total_shutt_dict_df[src][dest].at[item, current_date]
        for warehouse in warehouses:
            final_stock_dict_df[warehouse].at[item, current_date] = needs_shutt_stock_dict_df[warehouse].at[item, current_date]
            final_stock_hub_df.at[item, current_date] += final_stock_dict_df[warehouse].at[item, current_date]

    # ECCEDENZA IN NESSUN MAGAZZINO

    # CASO 1
    if (overcap_dict_df[6372].at['OVERCAP', current_date] == 0 and
            overcap_dict_df[6371].at['OVERCAP', current_date] == 0 and
            overcap_dict_df[6318].at['OVERCAP', current_date] == 0):
        no_overcap_count += 1
        print(f"In nessuno dei magazzini si supera la capacità alla data {current_date}")

    # ECCEDENZA IN 1 MAGAZZINO

    # CASO 2: 6372
    # Categoria item da poter spsotare: RUNNERS
    # Priorità nello spostamento: 72-71, 72,18
    # Crietrio: similmente a quanto fatto per lo scenario 4
    elif (overcap_dict_df[6372].at['OVERCAP', current_date] > 0 and
          overcap_dict_df[6371].at['OVERCAP', current_date] == 0 and
          overcap_dict_df[6318].at['OVERCAP', current_date] == 0):
        overcap_72_count += 1
        overcap72 = overcap_dict_df[6372].at['OVERCAP', current_date]
        print(f"Si supera la capacità solo in 72 alla data {current_date}, di una quantità pari a {overcap72}")
        # Calcolo lo spazio disponibile in 71 approssimandolo all' intero inferiore
        space71 = np.floor((warehouses_cap_dict['6371'] - needs_shutt_stock_dict_df[6371][current_date].sum()))
        space18 = np.floor((warehouses_cap_dict['6318'] - needs_shutt_stock_dict_df[6318][current_date].sum()))
        cap_shutt_72_71 = min(space71, overcap72)
        cap_shutt_72_18 = overcap72 - cap_shutt_72_71
        if overcap72 > space71 + space18:
            print(f"Attenzione: Non c' è spazio a sufficienza nell' hub alla data {current_date}.")
            exit(1)
        if overcap72 > space71 + space18:
            print(f"Attenzione: Non c' è spazio a sufficienza nell' hub alla data {current_date}.")
            exit(1)
        elif cap_shutt_72_71 > 0 and cap_shutt_72_18 == 0:
            print(f"Shuttling solo da 72 a 71 alla data {current_date}.")
            case2_only_cap_shutt_71_count += 1
        elif cap_shutt_72_71 == 0 and cap_shutt_72_18 > 0:
            print(f"Shuttling solo 72 a 18 alla data {current_date}.")
            case2_only_cap_shutt_18_count += 1
        elif cap_shutt_72_71 > 0 and cap_shutt_72_18 > 0:
            print(f"Shuttling sia 72 a 71 che da 72 a 18 alla data {current_date}.")
            case2_double_cap_shutt_count += 1
        # FASE 1: VERIFICA AMMISSIBILITA' SOGLIA DI BILANCIAMENTO
        # Troviamo la soglia di bilanciamento ammissibile attraverso un processo iterativo
        print("FASE 1: VERIFICA AMMISSIBILITA' SOGLIA DI BILANCIAMENTO")
        info_movable_items = {}
        # Categoria item da poter spsotare
        movable_items = runners
        # Definizione delle soglie iniziali
        balancing_thr_A = max_balancing_thr_A
        balancing_thr_AA = max_balancing_thr_AA
        while True:
            info_movable_items = {}
            alpha = None
            total_potential_shutt = 0
            for item in movable_items:
                # Tutti i runners hanno uscita media da entrambi i magazzini diversa da 0, dunque per ciascun item
                # salviamo l' uscitamedia di 71 e 18
                mean_out71 = mean_out_dict71.get(item, 0)
                mean_out18 = mean_out_dict18.get(item, 0)
                if mean_out71 == 0 and mean_out18 == 0:
                    continue  # Salta l'item corrente e passa al successivo
                # Determino la soglia di bilanciamento in base alla categoria dell'item
                if item in items_r_A:
                    alpha = balancing_thr_A
                elif item in items_r_AA:
                    alpha = balancing_thr_AA
                if next_date:
                    effective_stock72 = final_stock_dict_df[6372].at[item, current_date] - out_item_dict_df[6372].at[item, next_date]
                    effective_stock71 = final_stock_dict_df[6371].at[item, current_date] - out_item_dict_df[6371].at[item, next_date]
                    effective_stock18 = final_stock_dict_df[6318].at[item, current_date] - out_item_dict_df[6318].at[item, next_date]
                else:
                    effective_stock72 = final_stock_dict_df[6372].at[item, current_date]
                    effective_stock71 = final_stock_dict_df[6371].at[item, current_date]
                    effective_stock18 = final_stock_dict_df[6318].at[item, current_date]
                potential_shutt = max(0, effective_stock72 - np.ceil((effective_stock72 + effective_stock71 + effective_stock18) * alpha))
                # Inserisco le informazioni nel dizionario
                info_movable_items[item] = {
                    'mean_out71': mean_out71,
                    'mean_out18': mean_out18,
                    'effective_stock72': effective_stock72,
                    'effective_stock71': effective_stock71,
                    'effective_stock18': effective_stock18,
                    'potential_shutt': potential_shutt,
                }
            total_potential_shutt = sum(info['potential_shutt'] for info in info_movable_items.values())
            if total_potential_shutt >= overcap72:
                balancing_thr_df.at['soglia_finale_A', current_date] = balancing_thr_A
                balancing_thr_df.at['soglia_finale_AA', current_date] = balancing_thr_AA
                print(f"Soglia di bilanciamento ammissibile:")
                print(f" - AA: {balancing_thr_AA:.2f}")
                print(f" - A: {balancing_thr_A:.2f}")
                break
            elif balancing_thr_AA > min_balancing_thr_AA or balancing_thr_A > min_balancing_thr_A:
                balancing_thr_A = max(min_balancing_thr_A, balancing_thr_A - 0.02)
                balancing_thr_AA = max(min_balancing_thr_AA, balancing_thr_AA - 0.02)
                print(f"Soglia di bilanciamento AA abbassata a  {balancing_thr_AA:.2f}")
                print(f"Soglia di bilanciamento A abbassata a {balancing_thr_A:.2f}")
                continue  # Ripeti il ciclo con la nuova soglia abbassata
            else:
                print(f"ARRESTO FORZATO: la soglia di sicurezza è troppo alta e il while non può terminare.")
                print(f"Giorno di arresto: {current_date}")
                exit(1)
        # 72 - 71
        if cap_shutt_72_71 > 0 and cap_shutt_72_18 == 0:
            # FASE 2: SELEZIONE DEGLI ITEM E CALCOLO DEGLI SHUTTLING EFFETTIVI DA 72 A 71
            print(f"FASE 2: SELEZIONE DEGLI ITEM E CALCOLO DEGLI SHUTTLING EFFETTIVI DA 72 A 71 PER UN AMMONTARE PARI A {cap_shutt_72_71}")
            initial_cov_day_star = min((info['effective_stock71'] / mean_out_dict71[item] for item, info in info_movable_items.items()
                                        if mean_out_dict71.get(item, 0) > 0), default=0) + step_size  # Se nessun valore valido, restituisce step_size
            cov_day_star = initial_cov_day_star
            while True:
                # Creo una lista in cui salverò gli item da movimentare
                moved_items_dict = {}
                sum_total_eff_shutt = 0
                for item, info in info_movable_items.items():
                    cov_day = info['effective_stock71'] / info['mean_out71']
                    # Verifica se l'item è sotto la soglia di copertura attuale
                    if cov_day < cov_day_star:
                        cov_shutt = cov_day_star * info['mean_out71'] - info['effective_stock71']
                        eff_shutt = np.floor(min(cov_shutt, info['potential_shutt']))
                        if eff_shutt > 0:
                            moved_items_dict[item] = {
                                "effective_stock71": info['effective_stock71'],
                                "mean_out71": info['mean_out71'],
                                "cov_day": cov_day,
                                "eff_shutt": eff_shutt,
                            }
                            sum_total_eff_shutt += eff_shutt
                # Verifica se gli shuttling effettivi coprono cap_shutt_72_71 ed in tal caso esco dal while
                if sum_total_eff_shutt >= cap_shutt_72_71:
                    break
                # Incremento cov_day_star con step_size
                else:
                    cov_day_star += step_size
            difference = int(sum_total_eff_shutt - cap_shutt_72_71)  # Differenza da distribuire
            # Se difference > 0, dobbiamo sottrarre unità da eff_shutt in moved_items_dict
            if difference > 0:
                while difference > 0:
                    # Ordiniamo gli item in base a `eff_shutt` in ordine decrescente
                    sorted_items = sorted(moved_items_dict.keys(), key=lambda item: moved_items_dict[item]["eff_shutt"], reverse=True)
                    # Verifichiamo se abbiamo item disponibili per la riduzione
                    if not sorted_items:
                        print("Attenzione: Nessun item disponibile per ridurre la differenza.")
                        exit(1)
                    # Determiniamo quanti item possiamo aggiornare (massimo difference)
                    items_to_update = min(len(sorted_items), difference)
                    # Sottraiamo 1 ai primi items_to_update item per correggere la somma totale
                    for item in sorted_items[:items_to_update]:
                        moved_items_dict[item]["eff_shutt"] -= 1  # Riduciamo eff_shutt di 1
                        sum_total_eff_shutt -= 1  # Aggiorniamo la somma totale
                    # Aggiorniamo difference
                    difference -= items_to_update
            # Aggiorno il df degli shuttling effettivi totali
            total_eff_shutt_case2.at['6372-6371', current_date] = sum_total_eff_shutt
            # Ripartizione degli shuttling
            for item, data in moved_items_dict.items():
                # Aggiornamento dei DF
                cap_shutt_dict_df[6372][6371].at[item, current_date] = data["eff_shutt"]
                total_shutt_dict_df[6372][6371].at[item, current_date] += data["eff_shutt"]
                total_shutt_hub_df.at[item, current_date] += data["eff_shutt"]
                final_stock_dict_df[6372].at[item, current_date] -= data["eff_shutt"]
                final_stock_dict_df[6371].at[item, current_date] += data["eff_shutt"]

            total_cap_shutt_df.loc["6372-6371", current_date] = cap_shutt_dict_df[6372][6371][current_date].sum()
            # Converto il dizionario in un DataFrame
            # orient="index" indica che le chiavi principali del dizionario (ITEM_A, ITEM_B, ecc.) diventano l’indice delle righe del DataFrame
            # e che valori annidati (dizionari interni) vengono convertiti nelle colonne del DataFrame.
            # .reset_index() trasforma l’indice (ITEM_A, ITEM_B, ecc.) in una normale colonna.
            case2_results_df = pd.DataFrame.from_dict(moved_items_dict, orient="index").reset_index()
            case2_results_df.rename(columns={"index": "ITEM"}, inplace=True)
            # Aggiungo nomi colonna
            case2_results_df.columns = ['ITEM', 'STOCK NETTO', 'USCITA MEDIA', 'GIORNO DI COPERTURA INIZIALE','SHUTTLING EFFETTIVI']
            # Nome del foglio (sostituisce / con -)
            sheet_name = f"72_71_{current_date}_{cov_day_star:.2f}".replace('/', '-')
            with pd.ExcelWriter(output_step4_case2_path, mode='a', engine='openpyxl', if_sheet_exists='overlay') as writer:
                case2_results_df.to_excel(writer, sheet_name=sheet_name, index=False)

        # 72 - 18
        # FASE 2: SELEZIONE DEGLI ITEM E CALCOLO DEGLI SHUTTLING EFFETTIVI DA 72 A 18
        elif cap_shutt_72_71 == 0 and cap_shutt_72_18 > 0:
            print(f"FASE 2: SELEZIONE DEGLI ITEM E CALCOLO DEGLI SHUTTLING EFFETTIVI DA 72 A 18 PER UN AMMONTARE PARI A {cap_shutt_72_18}")
            initial_cov_day_star = min((info['effective_stock18'] / mean_out_dict18[item] for item, info in info_movable_items.items()
                                        if mean_out_dict18.get(item, 0) > 0), default=0) + step_size
            cov_day_star = initial_cov_day_star
            while True:
                moved_items_dict = {}
                sum_total_eff_shutt = 0
                for item, info in info_movable_items.items():
                    cov_day = info['effective_stock18'] / info['mean_out18']
                    if cov_day < cov_day_star:
                        cov_shutt = cov_day_star * info['mean_out18'] - info['effective_stock18']
                        eff_shutt = np.floor(min(cov_shutt, info['potential_shutt']))
                        if eff_shutt > 0:
                            moved_items_dict[item] = {
                                "effective_stock18": info['effective_stock18'],
                                "mean_out18": info['mean_out18'],
                                "cov_day": cov_day,
                                "eff_shutt": eff_shutt,
                            }
                            sum_total_eff_shutt += eff_shutt
                if sum_total_eff_shutt >= cap_shutt_72_18:
                    break
                else:
                    cov_day_star += step_size
            difference = int(sum_total_eff_shutt - cap_shutt_72_18)
            if difference > 0:
                while difference > 0:
                    sorted_items = sorted(moved_items_dict.keys(), key=lambda item: moved_items_dict[item]["eff_shutt"], reverse=True)
                    if not sorted_items:
                        print("Attenzione: Nessun item disponibile per ridurre la differenza.")
                        exit(1)
                    items_to_update = min(len(sorted_items), difference)
                    for item in sorted_items[:items_to_update]:
                        moved_items_dict[item]["eff_shutt"] -= 1
                        sum_total_eff_shutt -= 1
                    difference -= items_to_update
            total_eff_shutt_case2.at['6372-6318', current_date] = sum_total_eff_shutt
            for item, data in moved_items_dict.items():
                # Aggiornamento dei DF
                cap_shutt_dict_df[6372][6318].at[item, current_date] = data["eff_shutt"]
                total_shutt_dict_df[6372][6318].at[item, current_date] += data["eff_shutt"]
                total_shutt_hub_df.at[item, current_date] += data["eff_shutt"]
                final_stock_dict_df[6372].at[item, current_date] -= data["eff_shutt"]
                final_stock_dict_df[6318].at[item, current_date] += data["eff_shutt"]

            total_cap_shutt_df.loc["6372-6318", current_date] = cap_shutt_dict_df[6372][6318][current_date].sum()
            case2_results_df = pd.DataFrame.from_dict(moved_items_dict, orient="index").reset_index()
            case2_results_df.rename(columns={"index": "ITEM"}, inplace=True)
            case2_results_df.columns = ['ITEM', 'STOCK NETTO', 'USCITA MEDIA', 'GIORNO DI COPERTURA INIZIALE', 'SHUTTLING EFFETTIVI']
            sheet_name = f"72_18_{current_date}_{cov_day_star:.2f}".replace('/', '-')
            with pd.ExcelWriter(output_step4_case2_path, mode='a', engine='openpyxl', if_sheet_exists='overlay') as writer:
                case2_results_df.to_excel(writer, sheet_name=sheet_name, index=False)

        # 72 - 71 e 72 - 18
        # FASE 2: SELEZIONE DEGLI ITEM E CALCOLO DEGLI SHUTTLING EFFETTIVI DA 72 A 71 e 72 a 18
        elif cap_shutt_72_71 > 0 and cap_shutt_72_18 > 0:
            print(f"FASE 2: SELEZIONE DEGLI ITEM E CALCOLO DEGLI SHUTTLING EFFETTIVI DA 72 A 71 e 72 a 18 PER UN AMMONTARE PARI A {overcap72}")
            initial_cov_day_star = min(
                min((info['effective_stock71'] / mean_out_dict71[item]
                     for item, info in info_movable_items.items()
                     if mean_out_dict71.get(item, 0) > 0), default=0) + step_size,
                min((info['effective_stock18'] / mean_out_dict18[item]
                     for item, info in info_movable_items.items()
                     if mean_out_dict18.get(item, 0) > 0), default=0) + step_size
            )
            cov_day_star = initial_cov_day_star
            while True:
                moved_items_list71 = []
                moved_items_list18 = []
                sum_total_eff_shutt = 0
                remaining_space71 = cap_shutt_72_71
                remaining_space18 = cap_shutt_72_18
                for item, info in info_movable_items.items():
                    cov_day_71 = info['effective_stock71'] / info['mean_out71']
                    cov_day_18 = info['effective_stock18'] / info['mean_out18']
                    local_potential_shutt = info['potential_shutt']
                    if cov_day_18 < cov_day_star and cov_day_71 < cov_day_star:
                        if cov_day_18 <= cov_day_71 and remaining_space18 > 0:
                            cov_shutt = cov_day_star * info['mean_out18'] - info['effective_stock18']
                            eff_shutt = np.floor(min(cov_shutt, info['potential_shutt'], remaining_space18))
                            if eff_shutt > 0:
                                moved_items_list18.append({
                                    'item': item,
                                    'dest': '6318',
                                    'effective_stock': info['effective_stock18'],
                                    'mean_out': info['mean_out18'],
                                    'cov_day': cov_day_18,
                                    'eff_shutt': eff_shutt
                                })
                                sum_total_eff_shutt += eff_shutt
                                remaining_space18 -= eff_shutt
                                local_potential_shutt -= eff_shutt
                        if remaining_space71 > 0 and local_potential_shutt > 0:
                            cov_shutt = cov_day_star * info['mean_out71'] - info['effective_stock71']
                            eff_shutt = np.floor(min(cov_shutt, local_potential_shutt, remaining_space71))
                            if eff_shutt > 0:
                                moved_items_list71.append({
                                    'item': item,
                                    'dest': '6371',
                                    'effective_stock': info['effective_stock71'],
                                    'mean_out': info['mean_out71'],
                                    'cov_day': cov_day_71,
                                    'eff_shutt': eff_shutt
                                })
                                sum_total_eff_shutt += eff_shutt
                                remaining_space71 -= eff_shutt
                    elif cov_day_18 < cov_day_star and cov_day_71 >= cov_day_star:
                        if remaining_space18 > 0:
                            cov_shutt = cov_day_star * info['mean_out18'] - info['effective_stock18']
                            eff_shutt = np.floor(min(cov_shutt, info['potential_shutt'], remaining_space18))
                            if eff_shutt > 0:
                                moved_items_list18.append({
                                    'item': item,
                                    'dest': '6318',
                                    'effective_stock': info['effective_stock18'],
                                    'mean_out': info['mean_out18'],
                                    'cov_day': cov_day_18,
                                    'eff_shutt': eff_shutt
                                })
                                sum_total_eff_shutt += eff_shutt
                                remaining_space18 -= eff_shutt
                    elif cov_day_18 >= cov_day_star and cov_day_71 < cov_day_star:
                        cov_shutt = cov_day_star * info['mean_out71'] - info['effective_stock71']
                        eff_shutt = np.floor(min(cov_shutt, info['potential_shutt'], remaining_space71))
                        if eff_shutt > 0:
                            moved_items_list71.append({
                                'item': item,
                                'dest': '6371',
                                'effective_stock': info['effective_stock71'],
                                'mean_out': info['mean_out71'],
                                'cov_day': cov_day_71,
                                'eff_shutt': eff_shutt
                            })
                            sum_total_eff_shutt += eff_shutt
                            remaining_space71 -= eff_shutt
                if sum_total_eff_shutt >= overcap72:
                    break
                else:
                    cov_day_star += step_size
            # Ora unisco effettivamente tutti i movimenti
            moved_items_list = moved_items_list71 + moved_items_list18
            difference = int(sum_total_eff_shutt - overcap72)
            if difference > 0:
                difference_double_cap_shutt_count += 1
                # Unisco le due liste in un'unica lista
                # Ordino tutti gli item per eff_shutt decrescente
                moved_items_list_sorted = sorted(
                    moved_items_list,
                    key=lambda x: x['eff_shutt'],
                    reverse=True
                )
                # Inizio riduzione
                idx = 0
                while difference > 0 and idx < len(moved_items_list_sorted):
                    item_info = moved_items_list_sorted[idx]
                    if item_info['eff_shutt'] > 0:
                        item_info['eff_shutt'] -= 1
                        difference -= 1
                    # Se abbiamo ancora differenza, continua al prossimo item
                    idx = (idx + 1) % len(moved_items_list_sorted)
            total_eff_shutt_case2.at['6372-6371', current_date] = sum(item['eff_shutt'] for item in moved_items_list if item['dest'] == '6371')
            total_eff_shutt_case2.at['6372-6318', current_date] = sum(item['eff_shutt'] for item in moved_items_list if item['dest'] == '6318')

            # Ripartizione degli shuttling
            for data in moved_items_list:
                item = data['item']
                dest = int(data['dest'])
                eff_shutt = int(data['eff_shutt'])
                # Aggiornamento dei DataFrame per il magazzino di destinazione
                cap_shutt_dict_df[6372][dest].at[item, current_date] = eff_shutt
                total_shutt_dict_df[6372][dest].at[item, current_date] += eff_shutt
                total_shutt_hub_df.at[item, current_date] += eff_shutt
                final_stock_dict_df[6372].at[item, current_date] -= eff_shutt
                final_stock_dict_df[dest].at[item, current_date] += eff_shutt
            # Aggiornamento delle capacità totali di shuttling
            total_cap_shutt_df.loc['6372-6371', current_date] = cap_shutt_dict_df[6372][6371][current_date].sum()
            total_cap_shutt_df.loc['6372-6318', current_date] = cap_shutt_dict_df[6372][6318][current_date].sum()
            case2_results_df = pd.DataFrame(moved_items_list)
            case2_results_df.columns = ['ITEM', 'DEST', 'STOCK NETTO', 'USCITA MEDIA', 'GIORNO DI COPERTURA INIZIALE', 'SHUTTLING EFFETTIVI']
            # Nome del foglio (sostituisce / con -)
            sheet_name = f"72_71_18_{current_date}_{cov_day_star:.2f}".replace('/', '-')
            # Salvataggio nel file Excel
            with pd.ExcelWriter(output_step4_case2_path, mode='a', engine='openpyxl', if_sheet_exists='overlay') as writer:
                case2_results_df.to_excel(writer, sheet_name=sheet_name, index=False)

    # CASO 3: 6371
    elif (overcap_dict_df[6372].at['OVERCAP', current_date] == 0 and
          overcap_dict_df[6371].at['OVERCAP', current_date] > 0 and
          overcap_dict_df[6318].at['OVERCAP', current_date] == 0):
        overcap_71_count += 1

    # CASO 4: 6318
    elif (overcap_dict_df[6372].at['OVERCAP', current_date] == 0 and
          overcap_dict_df[6371].at['OVERCAP', current_date] == 0 and
          overcap_dict_df[6318].at['OVERCAP', current_date] > 0):
        overcap_18_count += 1

    # ECCEDENZA IN 2 MAGAZZINI

    # CASO 5: 6372, 6371
    elif (overcap_dict_df[6372].at['OVERCAP', current_date] > 0 and
          overcap_dict_df[6371].at['OVERCAP', current_date] > 0 and
          overcap_dict_df[6318].at['OVERCAP', current_date] == 0):
        overcap_72_71_count += 1

        # CASO 6: 6372, 6318
    elif (overcap_dict_df[6372].at['OVERCAP', current_date] > 0 and
          overcap_dict_df[6371].at['OVERCAP', current_date] == 0 and
          overcap_dict_df[6318].at['OVERCAP', current_date] > 0):
        overcap_72_18_count += 1

        # CASO 7: 6371, 6318
    elif (overcap_dict_df[6372].at['OVERCAP', current_date] == 0 and
          overcap_dict_df[6371].at['OVERCAP', current_date] > 0 and
          overcap_dict_df[6318].at['OVERCAP', current_date] > 0):
        overcap_71_18_count += 1

    # CASO 8: ECCEDENZA IN TUTTI I MAGAZZINI
    elif (overcap_dict_df[6372].at['OVERCAP', current_date] > 0 and
          overcap_dict_df[6371].at['OVERCAP', current_date] > 0 and
          overcap_dict_df[6318].at['OVERCAP', current_date] > 0):
        overcap_72_71_18_count += 1

    # Verifico che alla data corrente tutti e 3 i magazzini abbiano una scorta finale in grado di soddisfare
    # il rispettivo fabbisogno.
    for item in item_list:
        for warehouse in warehouses:
            if next_date:
                if final_stock_dict_df[warehouse].at[item, current_date] - out_item_dict_df[warehouse].at[
                    item, next_date] >= 0:
                    cap_check_passed += 1
                else:
                    cap_check_failed += 1
                    additional_stock = out_item_dict_df[warehouse].at[item, next_date] - final_stock_dict_df[warehouse].at[item, current_date]
                    print(f" Fabbisogno non correttamente soddisfatto alla data {current_date} per l' item {item}"
                          f" nel magazzino {warehouse} per una quantità pari a {additional_stock}, quindi correggo la scorta"
                          f" e aggiorno il dataframe della scorta aggiuntiva.")
                    final_stock_dict_df[warehouse].at[item, current_date] += additional_stock
                    additional_stock_df.at[item, "SCORTA AGGIUNTIVA"] += additional_stock
            else:
                if final_stock_dict_df[warehouse].at[item, current_date] >= 0:
                    cap_check_passed += 1
                else:
                    cap_check_failed += 1
                    additional_stock = - final_stock_dict_df[warehouse].at[item, current_date]
                    print(f" Fabbisogno non correttamente soddisfatto alla data {current_date} per l' item {item}"
                          f" nel magazzino {warehouse} per una quantità pari a {additional_stock}, quindi correggo la scorta"
                          f" e aggiorno il dataframe della scorta aggiuntiva.")
                    final_stock_dict_df[warehouse].at[item, current_date] += additional_stock
                    additional_stock_df.at[item, "SCORTA AGGIUNTIVA"] += additional_stock

    total_cap_shutt_df.loc["Totali"] = total_cap_shutt_df.drop(index="Totali").sum(axis=0)

# Filtro le sole date in cui si va in overcap di 72 (Caso 2 dello Step 4) per la creazione del df balancing_thr_df
overcap_date_72 = [date for date in date_list if overcap_dict_df[6372].at['OVERCAP', date] > 0 or
                   overcap_dict_df[6371].at['OVERCAP', date] > 0 or
                   overcap_dict_df[6318].at['OVERCAP', date] > 0]
filtered_balancing_thr_df = balancing_thr_df[overcap_date_72]

## INFO SHUTTLING ##

# SCOMPOSIZIONE TIPOLOGIE SHUTTLING TRA AA E A

def calculate_shuttling_totals(shuttling_dict, shuttling_type, warehouses):
    results = []
    for src in warehouses:
        for dest in warehouses:
            if src != dest:
                df = shuttling_dict[src][dest]
                total_r_AA = 0
                total_r_A = 0
                total_s_B = 0
                total_s_CC = 0
                total_s_C = 0
                for item in df.index:
                    if item in item_fam_dict:
                        fam = item_fam_dict[item]
                        if fam.endswith('AA'):
                            total_r_AA += df.loc[item].sum()
                        elif fam.endswith('A') and not fam.endswith('AA'):
                            total_r_A += df.loc[item].sum()
                        elif fam.endswith('B'):
                            total_s_B += df.loc[item].sum()
                        elif fam.endswith('CC'):
                            total_s_CC += df.loc[item].sum()
                        elif fam.endswith('C') and not fam.endswith('CC'):
                            total_s_C += df.loc[item].sum()

                # Aggiunge i risultati per il magazzino di uscita e ingresso
                results.append({
                    'Magazzino ORIGINE': src,
                    'Magazzino DESTINAZIONE': dest,
                    'Tipologia': shuttling_type,
                    'Categoria': 'AA',
                    'Totale': total_r_AA
                })
                results.append({
                    'Magazzino ORIGINE': src,
                    'Magazzino DESTINAZIONE': dest,
                    'Tipologia': shuttling_type,
                    'Categoria': 'A',
                    'Totale': total_r_A
                })
                results.append({
                    'Magazzino ORIGINE': src,
                    'Magazzino DESTINAZIONE': dest,
                    'Tipologia': shuttling_type,
                    'Categoria': 'B',
                    'Totale': total_s_B
                })
                results.append({
                    'Magazzino ORIGINE': src,
                    'Magazzino DESTINAZIONE': dest,
                    'Tipologia': shuttling_type,
                    'Categoria': 'CC',
                    'Totale': total_s_CC
                })
                results.append({
                    'Magazzino ORIGINE': src,
                    'Magazzino DESTINAZIONE': dest,
                    'Tipologia': shuttling_type,
                    'Categoria': 'C',
                    'Totale': total_s_C
                })
    return results
# Calcola gli shuttling totali per tipologia stratificati per categoria item
shuttling_results = []
shuttling_results.extend(calculate_shuttling_totals(needs_shutt_dict_df, 'FABB', warehouses))
shuttling_results.extend(calculate_shuttling_totals(cap_shutt_dict_df, 'CAP', warehouses))
# Creazione DataFrame per salvare i risultati
shuttling_summary_df = pd.DataFrame(shuttling_results)

# CALCOLO SHUTTLING TOTALI PER TIPOLOGIE

sum_total_shutt_type = {
    'SHUTT FABB 72-71': needs_shutt_dict_df[6372][6371].sum().sum(),
    'SHUTT FABB 72-18': needs_shutt_dict_df[6372][6318].sum().sum(),
    'SHUTT FABB 71-72': needs_shutt_dict_df[6371][6372].sum().sum(),
    'SHUTT FABB 71-18': round(needs_shutt_dict_df[6371][6318].sum().sum(), 2),
    'SHUTT FABB 18-72': needs_shutt_dict_df[6318][6372].sum().sum(),
    'SHUTT FABB 18-71': round(needs_shutt_dict_df[6318][6371].sum().sum(), 2),
    'SHUTT RITORNO': round((needs_shutt_dict_df[6371][6372].sum().sum() + needs_shutt_dict_df[6318][6372].sum().sum() + needs_shutt_dict_df[6318][6371].sum().sum()), 2),
    'SHUTT CAP 72-71': cap_shutt_dict_df[6372][6371].sum().sum(),
    'SHUTT CAP 72-18': cap_shutt_dict_df[6372][6318].sum().sum(),
    'SHUTT CAP 71-72': cap_shutt_dict_df[6371][6372].sum().sum(),
    'SHUTT CAP 71-18': round(cap_shutt_dict_df[6371][6318].sum().sum(), 2),
    'SHUTT CAP 18-72': cap_shutt_dict_df[6318][6372].sum().sum(),
    'SHUTT CAP 18-71': round(cap_shutt_dict_df[6318][6371].sum().sum(), 2),
    'SHUTT TOT 72-71': total_shutt_dict_df[6372][6371].sum().sum(),
    'SHUTT TOT 72-18': total_shutt_dict_df[6372][6318].sum().sum(),
    'SHUTT TOT 71-72': total_shutt_dict_df[6371][6372].sum().sum(),
    'SHUTT TOT 71-18': round(total_shutt_dict_df[6371][6318].sum().sum(), 2),
    'SHUTT TOT 18-72': total_shutt_dict_df[6318][6372].sum().sum(),
    'SHUTT TOT 18-71': round(total_shutt_dict_df[6318][6371].sum().sum(), 2),
    'SHUTT TOT HUB': round(total_shutt_hub_df.sum().sum(), 2)
}
# Creazione di un DataFrame per contenere i risultati
sum_total_shutt_type_df = pd.DataFrame(list(sum_total_shutt_type.items()), columns=['DESCRIZIONE', 'SOMMA'])

# SHUTTLING TOTALI PER FAMIGLIA E PER TIPOLOGIA

# Aggregazione dei totali per item per i tre DataFrame
needs_shutt_item_tot_72_71 = needs_shutt_dict_df[6372][6371].sum(axis=1)
needs_shutt_item_tot_72_18 = needs_shutt_dict_df[6372][6318].sum(axis=1)
needs_shutt_item_tot_71_72 = needs_shutt_dict_df[6371][6372].sum(axis=1)
needs_shutt_item_tot_71_18 = needs_shutt_dict_df[6371][6318].sum(axis=1)
needs_shutt_item_tot_18_72 = needs_shutt_dict_df[6318][6372].sum(axis=1)
needs_shutt_item_tot_18_71 = needs_shutt_dict_df[6318][6371].sum(axis=1)
cap_shutt_item_tot_72_71 = cap_shutt_dict_df[6372][6371].sum(axis=1)
cap_shutt_item_tot_72_18 = cap_shutt_dict_df[6372][6318].sum(axis=1)

# Creazione del dizionario per aggregare i totali per famiglia
shuttling_by_family = {}
# Scansioniamo ogni ITEM e sommiamo i valori per la famiglia corrispondente
for item, family in item_fam_dict.items():
    if family not in shuttling_by_family:
        shuttling_by_family[family] = {
            'SHUTT FABB 72-71': 0,
            'SHUTT FABB 72-18': 0,
            'SHUTT FABB 71-72': 0,
            'SHUTT FABB 71-18': 0,
            'SHUTT FABB 18-72': 0,
            'SHUTT FABB 18-71': 0,
            'SHUTT CAP 72-71': 0,
            'SHUTT CAP 72-18': 0,
        }
    # Se l'ITEM è presente nei DataFrame originali, sommiamo i totali alla sua famiglia
    if item in needs_shutt_item_tot_72_71.index:
        shuttling_by_family[family]["SHUTT FABB 72-71"] += needs_shutt_item_tot_72_71[item]
    if item in needs_shutt_item_tot_72_18.index:
        shuttling_by_family[family]["SHUTT FABB 72-18"] += needs_shutt_item_tot_72_18[item]
    if item in needs_shutt_item_tot_71_72.index:
        shuttling_by_family[family]["SHUTT FABB 71-72"] += needs_shutt_item_tot_71_72[item]
    if item in needs_shutt_item_tot_71_18.index:
        shuttling_by_family[family]["SHUTT FABB 71-18"] += needs_shutt_item_tot_71_18[item]
    if item in needs_shutt_item_tot_18_72.index:
        shuttling_by_family[family]["SHUTT FABB 18-72"] += needs_shutt_item_tot_18_72[item]
    if item in needs_shutt_item_tot_18_71.index:
        shuttling_by_family[family]["SHUTT FABB 18-71"] += needs_shutt_item_tot_18_71[item]
    if item in cap_shutt_item_tot_72_71.index:
        shuttling_by_family[family]["SHUTT CAP 72-71"] += cap_shutt_item_tot_72_71[item]
    if item in cap_shutt_item_tot_72_18.index:
        shuttling_by_family[family]["SHUTT CAP 72-18"] += cap_shutt_item_tot_72_18[item]

# Creiamo il DataFrame finale ordinato per famiglia
shuttling_by_family_df = pd.DataFrame.from_dict(shuttling_by_family, orient="index")
shuttling_by_family_df.index.name = "FAMIGLIA"

## STAMPE ##

print(f"Occorrenze overcap di 72: {overcap_72_count}")
sum_overcap72 = overcap_dict_df[6372].loc['OVERCAP'].sum() # Somma totale delle overcapacity di 6372
mean_overcap72 = sum_overcap72 / overcap_72_count
print(f"Overcap media di 72: {mean_overcap72:.0f}")
# Calcola la somma totale giornaliera delle scorte finali di 6372 (somma di tutte le righe per ogni giorno)
daily_total_stock72 = final_stock_dict_df[6372].sum(axis=0)
# Calcola la media dello stock totale giornaliero
mean_stock72 = daily_total_stock72.mean()
# Stampa il risultato
print(f"Livello di stock medio totale giornaliero di 6372: {mean_stock72:.0f}")
# Calcola la somma totale giornaliera delle scorte finali di 6371 (somma di tutte le righe per ogni giorno)
daily_total_stock71 = final_stock_dict_df[6371].sum(axis=0)
# Calcola la media dello stock totale giornaliero
mean_stock71 = daily_total_stock71.mean()
# Stampa il risultato
print(f"Livello di stock medio totale giornaliero di 6371: {mean_stock71:.0f}")
# Calcola la somma totale giornaliera delle scorte finali di 6318 (somma di tutte le righe per ogni giorno)
daily_total_stock18 = final_stock_dict_df[6318].sum(axis=0)
# Calcola la media dello stock totale giornaliero
mean_stock18 = daily_total_stock18.mean()
# Stampa il risultato
print(f"Livello di stock medio totale giornaliero di 6371: {mean_stock18:.0f}")

# Scorta media di runner A e AA per i magazzini
def calculate_mean_stock(final_stock_dict, items_r_AA, items_r_A, date_list):
    results = []
    for warehouse, df_scorte in final_stock_dict.items():
        # Totali giornalieri delle scorte per AA e A
        daily_AA_stock = df_scorte.loc[items_r_AA].sum(axis=0)  # Somma delle colonne per item AA
        daily_A_stock = df_scorte.loc[items_r_A].sum(axis=0)    # Somma delle colonne per item A
        # Media delle scorte giornaliere
        mean_AA_stock = daily_AA_stock.mean()
        mean_A_stock = daily_A_stock.mean()
        if warehouse in [6371, 6318]:
            daily_B_stock = df_scorte.loc[items_s_B].sum(axis=0)
            daily_CC_stock = df_scorte.loc[items_s_CC].sum(axis=0)
            daily_C_stock = df_scorte.loc[items_s_C].sum(axis=0)
            mean_B_stock = daily_B_stock.mean()
            mean_CC_stock = daily_CC_stock.mean()
            mean_C_stock = daily_C_stock.mean()
        else:
            mean_B_stock = 0
            mean_CC_stock = 0
            mean_C_stock = 0
        # Salva i results per il magazzino
        results.append({
            "Magazzino": warehouse,
            "Categoria": "AA",
            "Media Scorte": mean_AA_stock
        })
        results.append({
            "Magazzino": warehouse,
            "Categoria": "A",
            "Media Scorte": mean_A_stock
        })
        results.append({
            "Magazzino": warehouse,
            "Categoria": "B",
            "Media Scorte": mean_B_stock
        })
        results.append({
            "Magazzino": warehouse,
            "Categoria": "CC",
            "Media Scorte": mean_CC_stock
        })
        results.append({
            "Magazzino": warehouse,
            "Categoria": "C",
            "Media Scorte": mean_C_stock
        })
    return results

# Calcolo delle medie di scorte
mean_stock_results = calculate_mean_stock(final_stock_dict_df, items_r_AA, items_r_A, date_list)
# Creazione di un DataFrame per visualizzare e visualizzare i risultati
mean_stock_df = pd.DataFrame(mean_stock_results)
# Converte la colonna "Media Scorte" in interi
mean_stock_df["Media Scorte"] = mean_stock_df["Media Scorte"].astype(int)
# Stampa dei risultati in maniera ordinata
print("Livello medio di scorte per le categorie A e AA:")
print(mean_stock_df.to_string(index=False))
end_time = time.time()
execution_time = end_time - start_time
print(f"Tempo di esecuzione: {execution_time:.4f} secondi")

## COSTRUZIONE DOCUMENTO WORD ##

doc = Document()
# Aggiungi una sezione nel documento Word
doc.add_heading('Tempo di esecuzione', level=2)
doc.add_paragraph(f"Tempo di esecuzione totale: {execution_time:.4f} secondi")
doc.add_heading('Report Simulazione Scorte', level=1)
# Aggiungi i risultati
doc.add_heading('Overcapacity 6372', level=2)
doc.add_paragraph(f"Occorrenze overcap di 72: {overcap_72_count}")
doc.add_paragraph(f"Overcap media di 72: {mean_overcap72:.0f}")
doc.add_heading('Livello medio di stock totale giornaliero', level=2)
doc.add_paragraph(f"6372: {mean_stock72:.0f}")
doc.add_paragraph(f"6371: {mean_stock71:.0f}")
doc.add_paragraph(f"6318: {mean_stock18:.0f}")
doc.add_heading('Livello medio di scorte per categoria', level=2)
# Tabella
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Magazzino'
hdr_cells[1].text = 'Categoria'
hdr_cells[2].text = 'Media Scorte'
for index, row in mean_stock_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['Magazzino'])
    row_cells[1].text = row['Categoria']
    row_cells[2].text = str(row['Media Scorte'])
# Salvataggio
doc.save(output_word_path)
print(f"Report salvato in: {output_word_path}")

## SALVATAGGIO RISULTATI IN EXCEL ##

with pd.ExcelWriter(output_shipm_in_out_prod_path) as writer:
    for warehouse in warehouses:
        in_shipm_dict_df[warehouse].to_excel(writer, sheet_name=f'SPED_IN_{warehouse}', index=True)
        out_shipm_dict_df[warehouse].to_excel(writer, sheet_name=f'SPED_USCITA_{warehouse}', index=True)
        all_shipm_dict_df[warehouse].to_excel(writer, sheet_name=f'SPED_TOT_{warehouse}', index=True)
        in_item_dict_df[warehouse].to_excel(writer, sheet_name=f'INGRESSI_{warehouse}', index=True)
        out_item_dict_df[warehouse].to_excel(writer, sheet_name=f'USCITE_{warehouse}', index=True)
    complete_prod_df.to_excel(writer, sheet_name='PROD6372', index=False)

with pd.ExcelWriter(output_stock_path) as writer:
    for warehouse in warehouses:
        temporary_stock_dict_df[warehouse].to_excel(writer, sheet_name=f'SCORTE_PROV_{warehouse}', index=True)
        needs_shutt_stock_dict_df[warehouse].to_excel(writer, sheet_name=f'SCORTE_SH_FABB_{warehouse}', index=True)
        final_stock_dict_df[warehouse].to_excel(writer, sheet_name=f'SCORTE_FINALI_{warehouse}', index=True)
    final_stock_hub_df.to_excel(writer, sheet_name='SCORTE_FINALI_HUB', index=True)
    additional_stock_df.to_excel(writer, sheet_name='SCORTA_AGGIUNTIVA', index=True)

with pd.ExcelWriter(output_shuttling_and_cap_path) as writer:
    for src in warehouses:
        for dest in warehouses:
            if src != dest:
                needs_shutt_dict_df[src][dest].to_excel(writer, sheet_name=f'SH_FABB_{src}_{dest}', index=True)
                cap_shutt_dict_df[src][dest].to_excel(writer, sheet_name=f'SH_CAP_{src}_{dest}', index=True)
                total_shutt_dict_df[src][dest].to_excel(writer, sheet_name=f'SH_TOT_{src}_{dest}', index=True)
    for warehouse in warehouses:
        overcap_dict_df[warehouse].to_excel(writer, sheet_name=f'OVERCAP_{warehouse}', index=True)
    total_cap_shutt_df.to_excel(writer, sheet_name='SH_CAP_TOTALI', index=True)
    total_shutt_hub_df.to_excel(writer, sheet_name='SH_TOT_HUB', index=True)
    total_eff_shutt_case2.to_excel(writer, sheet_name='SH_TOT_EFF_CASO2', index=True)
    filtered_balancing_thr_df.to_excel(writer, sheet_name='SOGLIE_BIL')

with pd.ExcelWriter(output_shuttling_summary_path) as writer:
    sum_total_shutt_type_df.to_excel(writer, sheet_name='SHUTT_TIPO', index=False)
    shuttling_summary_df.to_excel(writer, sheet_name='STRATIFICAZIONE_AA_A', index=False)
    shuttling_by_family_df.to_excel(writer, sheet_name='SHUTT_FAM_TIPO')

print("I risultati sono stati esportati con successo nel file Excel di output!")









